# -*- coding: utf-8 -*-

"""
# reoredr_refs.py
This is an utility for reordering references in .docx manuscripts or reports.
If you don't use something like EndNote this is for you.

Usage:
python reorder_refs.py input.docx output.docx rules.conf

See the rules.conf sample for configuration details.

Written by prof. Artem Nikonorov artniko@gmail.com
"""

from docx import Document
import pytoml as toml
import argparse
import re
import sys
import pprint
from loguru import logger as log


# --------------------------------------------------------------------------
def logging_setup(config):
    log_file = sys.stdout

    format_long = (
        '<g>{time:YYYY-MM-DD hh:mm:ss}</g> | '
        '<level>{level:9}</level> | '
        '<c>{module}</c>:<c>{function}</c> - '
        '<b>{message}</b>'
    )
    format_short = (
        '<level>{level:9}</level> | '
        '<b>{message}</b>'
    )

    log.remove()
    log.add(
        log_file,
        format=format_short,
        level=config['options']['loglevel'],
    )


# --------------------------------------------------------------------------
def load_config(fname):
    with open(fname, 'rb') as f:
        #lines = f.readlines()
        config = toml.load(f)

    return config


# --------------------------------------------------------------------------
def load_rules(config):
    if 'reorder_rules' in config:
        rules = config['reorder_rules']
    else:
        pass

    old_refs = []
    new_refs = []

    for k in rules:
        old_bounds = k.split(":")
        old = list(range(int(old_bounds[0]), int(old_bounds[1]) + 1))
        new_bounds = rules[k].split(":")
        new = list(range(int(new_bounds[0]), int(new_bounds[1]) + 1))

        old_refs.extend(old)
        new_refs.extend(new)

    log.debug(old_refs)
    log.debug(new_refs)
    renumber_rules = {}
    ind = 0
    for k in old_refs:
        renumber_rules[k] = new_refs[ind]
        ind += 1

    return renumber_rules


# --------------------------------------------------------------------------
def check_refs(document, config):
    multi = False
    prefix = False

    for par in document.paragraphs:
        if '[' in par.text:
            multi_refs = re.findall(r'\[\D*\d+[ ]*-[ ]*\D*\d+]', par.text)
            if multi_refs:
                log.warning(f"Multi refs: {multi_refs}")
                multi = True

            prefix_refs = re.findall(r'\[\D+\d+\]', par.text)
            if prefix_refs:
                log.warning(f"Prefix refs: {prefix_refs}")
                prefix = True

    if multi:
        log.error(f"Multiple refs present, change these manually!\n")
    if prefix:
        log.error(f"Prefix refs present, change these manually!\n")

    if multi:
        if config['options'].get('stop_on_multiple_refs', True):
            raise Exception("Multiple refs")
    if prefix:
        if config['options'].get('stop_on_prefix_refs', True):
            raise Exception("Prefix refs")


# --------------------------------------------------------------------------
def save_reordered_refs(rules, long_ranges, config):
    file_name = config['options'].get('reordered_refs_file', "reordered_refs.txt")
    with open(file_name, 'w') as f:
        f.write("Reordering: old_ref -> new_ref\n")
        for old_ref in rules:
            f.write(f"{old_ref} -> {rules[old_ref]}\n")

        if long_ranges:
            f.write("\nFound long refs ranges:\n")
            f.write(f"{long_ranges}")


# --------------------------------------------------------------------------
def reorder_by_rules(document, rules, config):
    log.info("Reordering by rules started...")
    passed_refs = {}
    for old_ref, new_ref in rules.items():
        passed_refs[old_ref] = 0

    for par in document.paragraphs:
        if '[' in par.text:
            single_refs = re.findall(r'\[\d+]', par.text)
            if single_refs:
                log.debug(single_refs)
                new_text = par.text
                for ref in single_refs:
                    old_ref = int(ref[1:-1])

                    new_ref = rules[old_ref]
                    passed_refs[old_ref] += 1

                    log.info(f"{ref} -> [{new_ref}]")
                    new_text = new_text.replace(ref, f"[*{new_ref}]")
                    par.text = new_text

    errors = 0
    for old_ref, counter in passed_refs.items():
        if counter == 0:
            if errors == 0:
                log.error(f"Some old refs are missed!\n")

            errors += 1
            log.warning(f"Old ref [{old_ref}] never used!")

    save_reordered_refs(rules, [], config)

    return rules


# --------------------------------------------------------------------------
def auto_reorder(document, config):
    log.info("Auto reordering started...")

    rules = {}
    long_ranges = []
    for par in document.paragraphs:
        if '[' not in par.text:
            continue

        text = par.text
        old_ref_strings = re.findall(r'\[\D*\d+[ ]*-[ ]*\D*\d+]|\[\D*\d+]', text)
        log.debug(old_ref_strings)

        new_text = text
        for old_ref_string in old_ref_strings:
            r = old_ref_string[1:-1]
            r = r.replace(" ", "")
            ref_range = []
            if "-" in r:
                rr = r.split("-")
                prefs = re.findall(r'\D+', rr[0])
                pref = ""
                if prefs:
                    pref = prefs[0]
                    prefs = re.findall(r'\D+', rr[1])
                    if prefs[0] != pref:
                        log.error("Wrong prefix: {r}")
                        raise Exception("Wrong prefix")

                start = int(rr[0].replace(pref, ""))
                stop = int(rr[1].replace(pref, "")) + 1

                ref_range = list(range(start, stop))
                ref_range = list(map(lambda x: pref + str(x), ref_range))

            else:
                ref_range = [str(r)]

            new_ref_string = ""
            for ref in ref_range:
                if not rules.get(ref):
                    rules[ref] = len(rules) + 1

                new_ref_string += f"[*{rules[ref]}],"

            new_ref_string = new_ref_string[:-1]
            if len(ref_range) > 3:
                long_ranges.append(new_ref_string)

            new_text = new_text.replace(old_ref_string, new_ref_string)
            par.text = new_text

    str_rules = pprint.pformat(rules, sort_dicts=False)

    log.info(f"Reordered refs: \n{str_rules}")
    if long_ranges:
        log.warning("Found long refs ranges:")
        log.warning(long_ranges)

    save_reordered_refs(rules, long_ranges, config)

    return rules


# --------------------------------------------------------------------------
def reorder_ref_list(config, rules):
    ref_list_file = config["options"].get("ref_list_file")
    if not ref_list_file:
        log.info("\nReordering ref list skipped")
        return

    log.info("\nReordering ref list started...")
    try:
        document = Document(ref_list_file)
    except Exception as E:
        log.error(f"File not found: {ref_list_file}")
        log.error(f"Reordering ref list skipped")
        return

    old_list = []
    for par in document.paragraphs:
        old_list.append(par.text)

    ind = 0
    old_refs = list(rules)
    for par in document.paragraphs:
        if str(ind+1) not in old_refs:
            log.warning(f"Missed old ref [{ind+1}]")

        if ind >= len(old_refs):
            par.text = ""
            continue

        par.text = old_list[int(old_refs[ind]) - 1]
        ind += 1

    new_ref_list_file = ref_list_file.replace('.', '_reordered.')
    new_ref_list_file = config["options"].get("new_ref_list_file", new_ref_list_file)
    log.info(f"Saving reordered ref list to {new_ref_list_file}")
    document.save(new_ref_list_file)


# --------------------------------------------------------------------------
def parse_args():
    parser = argparse.ArgumentParser(
        description="References reordering in .docx documents",
        epilog="Example: python reorder_refs.py input.docx output.docx rules.conf")

    parser.add_argument('in_file', help='Input file name (.docx)')
    parser.add_argument('out_file', help='Output file name (.docx)')
    parser.add_argument('conf_file', help='Config file with reordering rules (.toml)')
    args = vars(parser.parse_args())
    return args


# --------------------------------------------------------------------------
def main():
    args = parse_args()

    config = load_config(args["conf_file"])

    logging_setup(config)

    log.info(f"Open input document {args['in_file']}")
    document = Document(args["in_file"])

    rules = []
    if config['options'].get('auto_reorder', False):
        rules = auto_reorder(document, config)
    else:
        rules = load_rules(config)
        check_refs(document, config)
        reorder_by_rules(document, rules, config)

    log.info(f"Saving output document to {args['out_file']}")
    document.save(args["out_file"])

    reorder_ref_list(config, rules)


# --------------------------------------------------------------------------
if __name__ == '__main__':
    main()
